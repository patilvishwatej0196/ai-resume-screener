# resume_reader.py
# Day 6 Update — Now supports BOTH PDF and DOCX resume formats
# Previously only supported PDF (Day 2)
# Now auto-detects file type and uses the correct reader

# Import pdfplumber — reads PDF files and extracts text
import pdfplumber

# Import Document from python-docx — reads Word document files
from docx import Document

# Import os — for file path and folder operations
import os


# -------------------------------------------------------
# FUNCTION 1: Extract text from a PDF file
# (same as Day 2 — no changes needed)
# -------------------------------------------------------

def extract_text_from_pdf(pdf_path):
    # Opens a PDF and extracts all text page by page
    # Returns the complete text as one string

    try:
        # Open the PDF using pdfplumber
        with pdfplumber.open(pdf_path) as pdf:

            # Empty list to collect text from each page
            all_text = []

            # Loop through every page in the PDF
            for page in pdf.pages:

                # Extract text from this page
                page_text = page.extract_text()

                # Only add if page had actual text
                # Some pages are image-only and return None
                if page_text:
                    all_text.append(page_text)

            # Join all pages with newline separator
            full_text = '\n'.join(all_text)

            return full_text

    except Exception as e:
        # Print error but don't crash
        print(f"Error reading PDF {pdf_path}: {e}")
        return ""


# -------------------------------------------------------
# FUNCTION 2: Extract text from a DOCX file (NEW - Day 6)
# -------------------------------------------------------

def extract_text_from_docx(docx_path):
    # Opens a Word document (.docx) and extracts all text
    # Word documents are structured as paragraphs
    # We loop through each paragraph and collect the text

    try:
        # Open the Word document using python-docx
        # Document() loads the entire .docx file into memory
        doc = Document(docx_path)

        # Empty list to collect text from each paragraph
        all_text = []

        # Loop through every paragraph in the document
        # A paragraph in Word = any block of text followed by Enter
        for paragraph in doc.paragraphs:

            # paragraph.text gives us the raw text of that paragraph
            text = paragraph.text

            # Only add non-empty paragraphs
            # strip() removes leading/trailing spaces
            if text.strip():
                all_text.append(text)

        # Also extract text from tables inside the document
        # Some resumes use tables for layout
        for table in doc.tables:

            # Loop through each row in the table
            for row in table.rows:

                # Loop through each cell in the row
                for cell in row.cells:

                    # Get text from this cell
                    cell_text = cell.text.strip()

                    # Only add non-empty cells
                    if cell_text:
                        all_text.append(cell_text)

        # Join all paragraphs with newline separator
        full_text = '\n'.join(all_text)

        return full_text

    except Exception as e:
        # Print error but don't crash the program
        print(f"Error reading DOCX {docx_path}: {e}")
        return ""


# -------------------------------------------------------
# FUNCTION 3: Auto-detect file type and extract text (NEW)
# -------------------------------------------------------

def extract_text(file_path):
    # This is the smart function that automatically detects
    # whether the file is PDF or DOCX and calls the right reader
    # This means the rest of the code never needs to check file type

    # Get the file extension in lowercase
    # os.path.splitext splits filename into (name, extension)
    # Example: "resume.PDF" -> (".pdf") -> ".pdf"
    extension = os.path.splitext(file_path)[1].lower()

    # If it's a PDF file — use pdfplumber reader
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)

    # If it's a Word document — use python-docx reader
    elif extension == '.docx':
        return extract_text_from_docx(file_path)

    # If it's an unknown file type — warn and return empty
    else:
        print(f"Unsupported file type: {extension}")
        print("Supported formats: .pdf and .docx")
        return ""


# -------------------------------------------------------
# FUNCTION 4: Read all resumes from a folder (UPDATED)
# -------------------------------------------------------

def read_all_resumes(folder_path):
    # Reads ALL resume files from a folder
    # Now supports both PDF and DOCX files
    # Returns a dictionary: { filename: extracted_text }

    # Empty dictionary to store results
    resumes = {}

    # Check the folder exists before trying to read it
    if not os.path.exists(folder_path):
        print(f"Folder not found: {folder_path}")
        return resumes

    # Define which file extensions we support
    # This makes it easy to add more formats later
    supported_extensions = ['.pdf', '.docx']

    # Loop through every file in the folder
    for filename in os.listdir(folder_path):

        # Get the file extension in lowercase
        extension = os.path.splitext(filename)[1].lower()

        # Only process supported file types
        if extension in supported_extensions:

            # Build the full path to this file
            full_path = os.path.join(folder_path, filename)

            # Show which file we are processing
            file_type = "PDF" if extension == '.pdf' else "DOCX"
            print(f"Reading [{file_type}]: {filename}")

            # Use auto-detect function to extract text
            text = extract_text(full_path)

            # Only save if we got actual text
            if text.strip():

                # Save to dictionary with filename as key
                resumes[filename] = text

                # Show how much text was extracted
                print(f"  Extracted {len(text)} characters")

            else:
                # Warn if file came back empty
                print(f"  Warning: No text found in {filename}")

    # Show summary
    print(f"\nTotal resumes loaded: {len(resumes)}")
    return resumes


# -------------------------------------------------------
# FUNCTION 5: Read a single job description text file
# (same as Day 2 — no changes needed)
# -------------------------------------------------------

def read_job_description(jd_path):
    # Reads a .txt job description file
    # Returns the content as a string

    try:
        # Open file in read mode with UTF-8 encoding
        with open(jd_path, 'r', encoding='utf-8') as file:
            content = file.read()
            return content

    except Exception as e:
        print(f"Error reading job description {jd_path}: {e}")
        return ""


# -------------------------------------------------------
# MAIN: Test everything when run directly
# -------------------------------------------------------

if __name__ == "__main__":

    print("=" * 55)
    print("RESUME READER - DAY 6 UPDATE TEST")
    print("Supports: PDF and DOCX formats")
    print("=" * 55)

    # ── TEST 1: Test DOCX reading ──
    print("\n[TEST 1] Testing DOCX text extraction...")
    print("-" * 40)

    # Create a sample DOCX file for testing
    # This saves a test Word document to your folder
    try:
        # Create a simple test Word document
        test_doc = Document()

        # Add content to the test document
        test_doc.add_heading('Test Candidate Name', 0)
        test_doc.add_paragraph('Email: testcandidate@gmail.com')
        test_doc.add_paragraph('Phone: 9876543210')
        test_doc.add_heading('Skills', level=1)
        test_doc.add_paragraph('Python, Machine Learning, Flask, SQL, Git, Docker')
        test_doc.add_heading('Experience', level=1)
        test_doc.add_paragraph('Data Science Intern at ABC Tech - June 2024')
        test_doc.add_paragraph('Built ML models using scikit-learn and pandas')

        # Save the test document
        test_path = "data/resumes/test_resume.docx"

        # Create the folder if it doesn't exist
        os.makedirs("data/resumes", exist_ok=True)

        # Save the file
        test_doc.save(test_path)
        print(f"Created test DOCX: {test_path}")

        # Now read it back using our new function
        extracted = extract_text_from_docx(test_path)

        if extracted:
            print(f"Successfully extracted {len(extracted)} characters")
            print("\nExtracted text preview:")
            print("-" * 40)
            print(extracted[:400])
        else:
            print("Failed to extract text from DOCX")

    except Exception as e:
        print(f"DOCX test failed: {e}")

    # ── TEST 2: Test auto-detect function ──
    print("\n\n[TEST 2] Testing auto-detect extract_text()...")
    print("-" * 40)

    # Test with the DOCX we just created
    test_files = [
        "data/resumes/test_resume.docx"
    ]

    # Add any PDFs if they exist
    resumes_folder = "data/resumes"
    if os.path.exists(resumes_folder):
        pdfs = [f for f in os.listdir(resumes_folder)
                if f.lower().endswith('.pdf')]
        for pdf in pdfs[:1]:  # test first PDF only
            test_files.append(os.path.join(resumes_folder, pdf))

    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\nAuto-detecting: {os.path.basename(file_path)}")
            text = extract_text(file_path)
            ext = os.path.splitext(file_path)[1].upper()
            print(f"File type: {ext}")
            print(f"Extracted: {len(text)} characters")
            print(f"Preview: {text[:100]}...")

    # ── TEST 3: Read all resumes from folder ──
    print("\n\n[TEST 3] Reading all resumes from data/resumes/...")
    print("-" * 40)

    all_resumes = read_all_resumes("data/resumes")

    print(f"\nSummary:")
    for filename, text in all_resumes.items():
        ext = os.path.splitext(filename)[1].upper()
        print(f"  [{ext}] {filename}: {len(text)} chars")

    print("\n" + "=" * 55)
    print("DAY 6 TEST COMPLETE")
    print("Both PDF and DOCX formats supported!")
    print("=" * 55)